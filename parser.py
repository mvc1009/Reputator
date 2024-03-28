import sys
from docx import *
from docx.shared import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

f = open(sys.argv[1], 'r')
json_object = json.load(f)

doc = Document("template.docx")
styles = doc.styles

# Create a custom styles for table cells
style = styles.add_style("Cell Text", WD_STYLE_TYPE.CHARACTER)
cell_text = doc.styles["Cell Text"]
cell_text_font = cell_text.font
cell_text_font.name = "Calibri"
cell_text_font.size = Pt(12)
cell_text_font.bold = True
cell_text_font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

style = styles.add_style("Cell Text Hit", WD_STYLE_TYPE.CHARACTER)
cell_text_hit = doc.styles["Cell Text Hit"]
cell_text_hit_font = cell_text_hit.font
cell_text_hit_font.name = "Calibri"
cell_text_hit_font.size = Pt(12)
cell_text_hit_font.bold = True
cell_text_hit_font.color.rgb = RGBColor(0x00, 0x96, 0x00)

style = styles.add_style("Cell Text Miss", WD_STYLE_TYPE.CHARACTER)
cell_text_miss = doc.styles["Cell Text Miss"]
cell_text_miss_font = cell_text_miss.font
cell_text_miss_font.name = "Calibri"
cell_text_miss_font.size = Pt(12)
cell_text_miss_font.bold = True
cell_text_miss_font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


table = doc.add_table(rows=1, cols=3, style='GoReport')
table.autofit = True
table.allow_autofit = True

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Scope / IP'
hdr_cells[1].text = 'Blacklisted'
hdr_cells[2].text = 'Lists'

for result in json_object:
	print("[+] %s" % result['ip'])
	listed = result['resume'].split(' ')[1]
	row_cells = table.add_row().cells
	row_cells[0].text = result['ip']
	
	if listed == '0':
		row_cells[1].paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
		row_cells[2].text = '-'
	else:
		row_cells[1].paragraphs[0].add_run(u'\u2718', "Cell Text Miss")
		blacklists = ''
		for i in result['results']:
			if i[0] == 'LISTED':
				blacklists += i[1] + ', '
		row_cells[2].text = blacklists[:-2]


doc_para = doc.add_paragraph('')
doc.save('./out.docx')

